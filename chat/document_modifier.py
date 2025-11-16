# FICHIER: chat/document_modifier.py
# Service pour modifier des documents en préservant leur mise en forme originale

import os
import re
from io import BytesIO
from pathlib import Path
from typing import Dict, Tuple, Optional
from django.conf import settings

# Imports conditionnels
try:
    from docx import Document as DocxDocument
    from docx.shared import RGBColor, Pt
    from docx.enum.text import WD_COLOR_INDEX
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from groq import Groq
    GROQ_AVAILABLE = True
except ImportError:
    GROQ_AVAILABLE = False


class DocumentModifierService:
    """
    Service pour modifier des documents en préservant leur structure originale
    (tableaux, colonnes, styles, mise en page)
    """

    @staticmethod
    def apply_changes_to_file(original_file_path: str, doc1_content: str, doc2_content: str,
                               doc1_title: str, doc2_title: str, pdf_structure_stored: dict = None) -> Tuple[bool, str, Optional[BytesIO]]:
        """
        Applique les changements à un fichier en préservant sa structure

        Args:
            original_file_path: Chemin du fichier original
            doc1_content: Contenu texte du document 1
            doc2_content: Contenu texte du document 2
            doc1_title: Titre du document 1
            doc2_title: Titre du document 2
            pdf_structure_stored: Structure PDF extraite lors de l'upload (optionnel)

        Returns:
            Tuple (success, message, modified_file_buffer)
        """
        file_ext = Path(original_file_path).suffix.lower()

        if file_ext == '.docx':
            return DocumentModifierService._modify_docx(
                original_file_path, doc1_content, doc2_content, doc1_title, doc2_title
            )
        elif file_ext == '.pdf':
            return DocumentModifierService._modify_pdf(
                original_file_path, doc1_content, doc2_content, doc1_title, doc2_title, pdf_structure_stored
            )
        else:
            return False, f"Format de fichier non supporté: {file_ext}", None

    @staticmethod
    def _modify_docx(file_path: str, doc1_content: str, doc2_content: str,
                     doc1_title: str, doc2_title: str) -> Tuple[bool, str, Optional[BytesIO]]:
        """
        Modifie un fichier Word (.docx) en préservant sa structure
        """
        if not DOCX_AVAILABLE:
            return False, "python-docx n'est pas installé", None

        try:
            # Obtenir les modifications à faire du LLM
            modifications = DocumentModifierService._get_modifications_from_llm(
                doc1_content, doc2_content, doc1_title, doc2_title
            )

            if not modifications:
                return False, "Impossible d'obtenir les modifications du LLM", None

            # Charger le document Word
            doc = DocxDocument(file_path)

            # Appliquer les modifications
            changes_made = 0

            # Parcourir tous les paragraphes
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    original_text = paragraph.text
                    modified_text = DocumentModifierService._apply_text_modifications(
                        original_text, modifications
                    )
                    if modified_text != original_text:
                        # Remplacer le texte en préservant le style
                        DocumentModifierService._replace_paragraph_text(paragraph, modified_text)
                        changes_made += 1

            # Parcourir tous les tableaux
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if paragraph.text.strip():
                                original_text = paragraph.text
                                modified_text = DocumentModifierService._apply_text_modifications(
                                    original_text, modifications
                                )
                                if modified_text != original_text:
                                    DocumentModifierService._replace_paragraph_text(paragraph, modified_text)
                                    changes_made += 1

            # Sauvegarder dans un buffer
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            message = f"Document modifié avec succès. {changes_made} éléments mis à jour."
            return True, message, buffer

        except Exception as e:
            return False, f"Erreur lors de la modification du document Word: {str(e)}", None

    @staticmethod
    def _modify_pdf(file_path: str, doc1_content: str, doc2_content: str,
                    doc1_title: str, doc2_title: str, pdf_structure_stored: dict = None) -> Tuple[bool, str, Optional[BytesIO]]:
        """
        Pour les PDF, on utilise la structure stockée (extraite lors de l'upload) et on génère un nouveau PDF
        en préservant autant que possible la structure originale
        """
        from .pdf_generator import PDFDocumentGenerator
        from .pdf_extractor import PDFStructureExtractor

        try:
            # Utiliser la structure stockée si disponible, sinon extraire
            if pdf_structure_stored and pdf_structure_stored.get('success'):
                print(f"[INFO] Utilisation de la structure PDF stockée ({pdf_structure_stored.get('total_tables', 0)} tableau(x))")
                pdf_structure = pdf_structure_stored
            else:
                print("[INFO] Extraction de la structure PDF à la volée")
                # Extraire la structure du PDF original (avec tableaux)
                pdf_structure = PDFStructureExtractor.extract_structure(file_path)

            if not pdf_structure or not pdf_structure.get('success'):
                # Fallback : génération simple sans structure
                return DocumentModifierService._modify_pdf_simple(
                    file_path, doc1_content, doc2_content, doc1_title, doc2_title
                )

            # Obtenir les modifications à faire du LLM
            modifications = DocumentModifierService._get_modifications_from_llm(
                doc1_content, doc2_content, doc1_title, doc2_title
            )

            if not modifications:
                return False, "Impossible d'obtenir les modifications du LLM", None

            # Générer un nouveau PDF avec la structure préservée
            pdf_generator = PDFDocumentGenerator()
            buffer = pdf_generator.generate_pdf_from_structure(
                title=doc1_title,
                pdf_structure=pdf_structure,
                modifications=modifications
            )

            tables_count = sum(len(page.get('tables', [])) for page in pdf_structure.get('pages', []))
            message = f"PDF généré avec structure préservée ({tables_count} tableau(x) recréé(s))"
            return True, message, buffer

        except Exception as e:
            print(f"[ERROR] _modify_pdf: {e}")
            # Fallback en cas d'erreur
            return DocumentModifierService._modify_pdf_simple(
                file_path, doc1_content, doc2_content, doc1_title, doc2_title
            )

    @staticmethod
    def _modify_pdf_simple(file_path: str, doc1_content: str, doc2_content: str,
                           doc1_title: str, doc2_title: str) -> Tuple[bool, str, Optional[BytesIO]]:
        """
        Fallback : Génération simple de PDF sans préservation de structure
        """
        from .pdf_generator import PDFDocumentGenerator

        try:
            # Obtenir le contenu fusionné du LLM
            merged_content = DocumentModifierService._get_merged_content_from_llm(
                doc1_content, doc2_content, doc1_title, doc2_title
            )

            if not merged_content:
                return False, "Impossible d'obtenir le contenu fusionné du LLM", None

            # Générer un nouveau PDF
            pdf_generator = PDFDocumentGenerator()
            buffer = pdf_generator.generate_clean_document_pdf(
                title=doc1_title,
                content=merged_content
            )

            message = "PDF généré avec le contenu mis à jour (structure simple)"
            return True, message, buffer

        except Exception as e:
            return False, f"Erreur lors de la génération du PDF: {str(e)}", None

    @staticmethod
    def _get_modifications_from_llm(doc1_content: str, doc2_content: str,
                                     doc1_title: str, doc2_title: str) -> Dict:
        """
        Utilise le LLM pour obtenir une liste de modifications à appliquer
        """
        if not GROQ_AVAILABLE or not hasattr(settings, 'GROQ_API_KEY') or not settings.GROQ_API_KEY:
            return {}

        try:
            print("[INFO] Demande de modifications au LLM...")
            client = Groq(api_key=settings.GROQ_API_KEY)

            # Limiter la longueur
            max_length = 4000
            content1_truncated = doc1_content[:max_length]
            content2_truncated = doc2_content[:max_length]

            print(f"[INFO] Longueurs: Doc1={len(content1_truncated)} chars, Doc2={len(content2_truncated)} chars")

            prompt = f"""Tu es un expert en édition de documents. Analyse ces deux documents et fournis une liste de modifications précises.

DOCUMENT 1 (à modifier): "{doc1_title}"
{content1_truncated}

DOCUMENT 2 (source des améliorations): "{doc2_title}"
{content2_truncated}

TÂCHE:
Identifie les différences entre les deux documents et fournis une liste de modifications sous forme de paires AVANT → APRÈS.

RÈGLES IMPORTANTES:
1. Donne des modifications PRÉCISES et COURTES (phrases ou paragraphes spécifiques)
2. Ne modifie PAS les titres de sections, les en-têtes ou la structure
3. Focus sur le contenu textuel qui a changé
4. Fournis maximum 15 modifications les plus importantes
5. Format EXACT pour chaque modification:
   AVANT: [texte exact à trouver]
   APRÈS: [texte de remplacement]
   ---

Fournis les modifications maintenant:"""

            response = client.chat.completions.create(
                model=settings.GROQ_MODEL,
                messages=[
                    {
                        "role": "system",
                        "content": "Tu es un assistant expert en édition de documents. Tu fournis des modifications précises et ciblées."
                    },
                    {
                        "role": "user",
                        "content": prompt
                    }
                ],
                temperature=0.2,
                max_tokens=3000
            )

            modifications_text = response.choices[0].message.content

            print(f"[INFO] Réponse LLM reçue: {len(modifications_text)} caractères")
            print(f"[DEBUG] Extrait de la réponse: {modifications_text[:200]}...")

            # Parser les modifications
            modifications = DocumentModifierService._parse_modifications(modifications_text)

            print(f"[INFO] {len(modifications)} modifications parsées")

            if not modifications:
                print("[WARNING] Aucune modification parsée depuis la réponse LLM !")
                print(f"[DEBUG] Réponse complète: {modifications_text}")

            return modifications

        except Exception as e:
            print(f"[ERROR] Erreur LLM modifications: {type(e).__name__}: {e}")
            import traceback
            traceback.print_exc()
            return {}

    @staticmethod
    def _parse_modifications(modifications_text: str) -> Dict:
        """
        Parse le texte de modifications pour extraire les paires AVANT → APRÈS
        """
        modifications = {}

        # Pattern plus flexible : accepte les numéros avant AVANT:
        # Exemple: "1. AVANT : texte" ou "AVANT: texte" ou "AVANT : texte"
        pattern = r'(?:\d+\.\s*)?AVANT\s*:\s*(.+?)\s*APRÈS\s*:\s*(.+?)\s*---'
        matches = re.findall(pattern, modifications_text, re.DOTALL | re.IGNORECASE)

        print(f"[DEBUG] Pattern regex trouvé {len(matches)} correspondances")

        for i, (avant, apres) in enumerate(matches):
            avant_clean = avant.strip().strip('[]"\'')
            apres_clean = apres.strip().strip('[]"\'')

            # Enlever les commentaires entre parenthèses
            apres_clean = re.sub(r'\s*\(pas de changement.*?\)', '', apres_clean, flags=re.IGNORECASE)
            apres_clean = apres_clean.strip()

            if avant_clean and apres_clean and avant_clean != apres_clean:
                modifications[avant_clean] = apres_clean
                print(f"[DEBUG] Modification {i+1}: '{avant_clean[:50]}...' → '{apres_clean[:50]}...'")

        return modifications

    @staticmethod
    def _apply_text_modifications(text: str, modifications: Dict) -> str:
        """
        Applique les modifications à un texte donné
        """
        modified_text = text

        for old_text, new_text in modifications.items():
            # Recherche flexible (ignorer la casse et les espaces multiples)
            if old_text.lower() in modified_text.lower():
                # Trouver la position exacte
                pattern = re.escape(old_text)
                modified_text = re.sub(pattern, new_text, modified_text, flags=re.IGNORECASE)

        return modified_text

    @staticmethod
    def _replace_paragraph_text(paragraph, new_text: str):
        """
        Remplace le texte d'un paragraphe en préservant le formatage
        """
        # Sauvegarder le style du premier run
        original_style = None
        if paragraph.runs:
            first_run = paragraph.runs[0]
            original_style = {
                'bold': first_run.bold,
                'italic': first_run.italic,
                'underline': first_run.underline,
                'font_name': first_run.font.name,
                'font_size': first_run.font.size,
            }

        # Supprimer tous les runs existants
        for run in paragraph.runs:
            run.text = ''

        # Créer un nouveau run avec le nouveau texte
        new_run = paragraph.add_run(new_text)

        # Appliquer le style original si disponible
        if original_style:
            if original_style['bold'] is not None:
                new_run.bold = original_style['bold']
            if original_style['italic'] is not None:
                new_run.italic = original_style['italic']
            if original_style['underline'] is not None:
                new_run.underline = original_style['underline']
            if original_style['font_name']:
                new_run.font.name = original_style['font_name']
            if original_style['font_size']:
                new_run.font.size = original_style['font_size']

    @staticmethod
    def _get_merged_content_from_llm(doc1_content: str, doc2_content: str,
                                      doc1_title: str, doc2_title: str) -> str:
        """
        Obtient le contenu fusionné du LLM pour les PDF
        """
        if not GROQ_AVAILABLE or not hasattr(settings, 'GROQ_API_KEY') or not settings.GROQ_API_KEY:
            return doc1_content

        try:
            client = Groq(api_key=settings.GROQ_API_KEY)

            # Limiter la longueur
            max_length = 5000
            content1_truncated = doc1_content[:max_length]
            content2_truncated = doc2_content[:max_length]

            prompt = f"""Tu es un expert en fusion de documents. Crée une version mise à jour du Document 1 en intégrant les améliorations du Document 2.

DOCUMENT 1 (base): "{doc1_title}"
{content1_truncated}

DOCUMENT 2 (améliorations): "{doc2_title}"
{content2_truncated}

INSTRUCTIONS:
1. Prends le Document 1 comme structure de base
2. Intègre les améliorations, corrections et ajouts du Document 2
3. NE METS AUCUN marqueur comme [MODIFIÉ] ou [AJOUTÉ]
4. Conserve la structure originale (titres, sections, paragraphes)
5. Le résultat doit être naturel et fluide
6. Si le Document 1 a des tableaux ou listes, essaie de les indiquer avec du formatage texte approprié

Génère le document fusionné maintenant:"""

            response = client.chat.completions.create(
                model=settings.GROQ_MODEL,
                messages=[
                    {
                        "role": "system",
                        "content": "Tu es un assistant expert en fusion de documents. Tu produis des documents propres et cohérents."
                    },
                    {
                        "role": "user",
                        "content": prompt
                    }
                ],
                temperature=0.3,
                max_tokens=4000
            )

            merged_content = response.choices[0].message.content

            # Nettoyer les marqueurs éventuels
            merged_content = merged_content.replace('[MODIFIÉ]', '').replace('[AJOUTÉ]', '')
            merged_content = merged_content.replace('[MODIFIED]', '').replace('[ADDED]', '')

            return merged_content.strip()

        except Exception as e:
            print(f"[ERROR] Erreur LLM fusion: {e}")
            return doc1_content
